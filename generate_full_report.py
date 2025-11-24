import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    # Initialize Document
    doc = Document()

    # ---------------------------------------------------------
    # SETUP: STYLES & MARGINS
    # ---------------------------------------------------------
    # Set margins (approx 1 inch = 2.54cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Enable different first page header/footer (to hide header on cover)
        section.different_first_page_header_footer = True

    # Define a standard font helper
    def add_formatted_text(paragraph, text, font_name='Arial', size=11, bold=False, color=None, italic=False):
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return run

    # ---------------------------------------------------------
    # HEADER & FOOTER (Applies to Page 2 onwards)
    # ---------------------------------------------------------
    # Access the default header/footer (for pages 2+)
    header = sections[0].header
    htext = header.paragraphs[0]
    htext.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_text(htext, "VITyarthi Project: SpamGuard AI", size=8, color=RGBColor(128, 128, 128), italic=True)

    footer = sections[0].footer
    ftext = footer.paragraphs[0]
    ftext.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Note: Word handles page numbers dynamically, simpler to just put text for now
    add_formatted_text(ftext, "SpamGuard AI Report", size=8, color=RGBColor(128, 128, 128), italic=True)

    # ---------------------------------------------------------
    # PAGE 1: COVER PAGE
    # ---------------------------------------------------------
    # Add spacing at top
    for _ in range(4): doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "PROJECT REPORT", size=24, bold=True)

    # Decorative Line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("__________________________________________________")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph() # Spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "SpamGuard AI", size=28, bold=True, color=RGBColor(0, 51, 102))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "SMS & Email Phishing Classifier", size=16)

    for _ in range(5): doc.add_paragraph() # Spacing

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "Submitted by:", size=14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "[YOUR NAME HERE]", size=14, bold=True) # <--- EDIT THIS

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "[YOUR REG NO HERE]", size=14, bold=True) # <--- EDIT THIS

    for _ in range(3): doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "Course: Programming for Problem Solving\nSchool of Computer Science and Engineering\nNovember 2025", size=12)

    doc.add_page_break()

    # ---------------------------------------------------------
    # HELPER FUNCTIONS FOR BODY
    # ---------------------------------------------------------
    def add_chapter_title(text):
        p = doc.add_heading(text, level=1)
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14)

    def add_body_text(text):
        p = doc.add_paragraph(text)
        p.style = doc.styles['Normal']
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.size = Pt(11)

    def add_bullet(text):
        p = doc.add_paragraph(text, style='List Bullet')
        if p.runs:
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(11)

    def add_image_safe(path, caption):
        if os.path.exists(path):
            try:
                doc.add_picture(path, width=Inches(6))
                # Caption
                c = doc.add_paragraph(caption)
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.runs[0].font.italic = True
                c.runs[0].font.size = Pt(9)
            except Exception as e:
                print(f"Error adding image {path}: {e}")
        else:
            p = doc.add_paragraph(f"[Image Placeholder: {path} - Missing File]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

    # ---------------------------------------------------------
    # CONTENT SECTIONS
    # ---------------------------------------------------------

    # 1. Introduction
    add_chapter_title("1. Introduction")
    add_body_text("In the modern digital landscape, Short Message Service (SMS) and Email have become the primary communication channels for personal and professional interaction. However, this ubiquity has made them a prime target for cybercriminals. 'Smishing' (SMS Phishing) and spam campaigns are used to deceive users into revealing sensitive information or installing malware.")
    add_body_text("SpamGuard AI is a Machine Learning (ML) based utility designed to automatically detect and filter these malicious messages. Unlike traditional rule-based filters that rely on static keyword lists, SpamGuard utilizes Natural Language Processing (NLP) and a Naive Bayes classifier to learn patterns from data, enabling it to distinguish between legitimate ('Ham') and malicious ('Spam') communications with high accuracy.")

    # 2. Problem Statement
    add_chapter_title("2. Problem Statement")
    
    p = doc.add_paragraph()
    add_formatted_text(p, "The Core Problem: ", bold=True)
    p.add_run("Users are inundated with unsolicited messages ranging from marketing spam to dangerous phishing attempts.")

    p = doc.add_paragraph()
    add_formatted_text(p, "The Technical Gap: ", bold=True)
    p.add_run("SMS often lacks robust, client-side filtering. There is a need for a lightweight, trainable system that can classify messages offline.")

    p = doc.add_paragraph()
    add_formatted_text(p, "Solution: ", bold=True)
    p.add_run("SpamGuard AI addresses this by providing a local, command-line interface tool that pre-processes text and predicts the likelihood of spam.")

    # 3. Functional Requirements
    add_chapter_title("3. Functional Requirements")
    add_body_text("The system is built upon three distinct functional modules:")
    add_bullet("Module 1: Data Preprocessing. Removes punctuation, lowers case, tokenizes text.")
    add_bullet("Module 2: Model Training. Trains a Multinomial Naive Bayes classifier.")
    add_bullet("Module 3: Prediction Interface. Menu-driven CLI for real-time feedback.")

    # 4. Non-Functional Requirements
    add_chapter_title("4. Non-Functional Requirements")
    add_bullet("Accuracy: Minimum 85% on standard test datasets.")
    add_bullet("Performance: Prediction < 1 second.")
    add_bullet("Reliability: Model persistence via disk storage.")
    add_bullet("Usability: Graceful exception handling.")

    # 5. Architecture
    doc.add_page_break()
    add_chapter_title("5. System Architecture")
    add_body_text("The project follows a modular Machine Learning Pipeline architecture: Data Layer, Preprocessing Layer, Logic Layer (Model), and Presentation Layer.")
    add_image_safe("System_Architecture.png", "Figure 1: System Architecture Diagram")

    # 6. Diagrams
    doc.add_page_break()
    add_chapter_title("6. Design Diagrams")
    
    p = doc.add_paragraph("6.1 Workflow Diagram")
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(12)
    add_body_text("The flow illustrates the user decision process: Training -> Saving -> Predicting.")
    add_image_safe("Workflow_Diagram.png", "Figure 2: Workflow / Process Flow")
    
    doc.add_paragraph() # spacer

    p = doc.add_paragraph("6.2 Use Case Diagram")
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(12)
    add_body_text("Primary actor (User) interacts with Train, Input Message, View Prediction, and Exit.")
    add_image_safe("Use_Case_Diagram.png", "Figure 3: Use Case Diagram")

    # 7. Design Decisions
    doc.add_page_break()
    add_chapter_title("7. Design Decisions & Rationale")
    add_bullet("Why Naive Bayes? Speed and high performance on text data (Bag of Words).")
    add_bullet("Why Scikit-Learn? Robust pipelines and consistency.")
    add_bullet("Storage Format: Joblib for efficient NumPy array storage.")

    # 8. Implementation
    add_chapter_title("8. Implementation Details")
    add_bullet("preprocessing.py: Text cleaning.")
    add_bullet("model_trainer.py: Training logic.")
    add_bullet("predictor.py: Inference engine.")
    add_bullet("main.py: User flow orchestration.")

    # 9. Screenshots
    doc.add_page_break()
    add_chapter_title("9. Screenshots")
    add_image_safe("screenshot_menu.png", "Figure 4: Main Menu Interface")
    add_image_safe("screenshot_output.png", "Figure 5: Prediction Result")

    # 10. Testing
    add_chapter_title("10. Testing Approach")
    add_bullet("Positive Testing: Validating 'Free Lottery' as Spam.")
    add_bullet("Boundary Testing: Empty strings/numbers.")
    add_bullet("Persistence Testing: Restarting app to check model load.")

    # 11. Conclusion
    add_chapter_title("11. Challenges & Future")
    add_body_text("Challenges: Handling case sensitivity and cold starts. Learnings: NLP Tokenization and modular coding.")
    add_bullet("Future: Web Interface (Streamlit).")
    add_bullet("Future: Deep Learning (BERT).")

    # 12. References
    add_chapter_title("12. References")
    add_bullet("Scikit-Learn Documentation")
    add_bullet("Pandas Documentation")

    # Save
    try:
        filename = "SpamGuard_Report_Editable.docx"
        doc.save(filename)
        print(f"SUCCESS! Editable Word document generated: {filename}")
    except Exception as e:
        print(f"Error saving document: {e}")
        print("Make sure the file isn't already open in Word!")

if __name__ == "__main__":
    create_report()